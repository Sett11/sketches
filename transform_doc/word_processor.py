#!/usr/bin/env python3
"""
Простой процессор Word документов на python-docx
Тестирует сохранение форматирования при изменении содержимого
"""

import os
import traceback
from datetime import datetime
from typing import Dict, List
from dotenv import load_dotenv

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from mylogger import Logger
from llm import OpenRouterClient

# Загружаем переменные окружения из .env файла
load_dotenv()

class WordProcessor:
    """Простой процессор Word документов с сохранением форматирования"""
    
    def __init__(self):
        self.document = None
        self.logger = Logger('WordProcessor', 'logs/word_processor.log')
        self.logger.info("WordProcessor инициализирован")
    
    def load_document(self, filepath: str) -> 'WordProcessor':
        """Загружает документ из файла"""
        self.logger.info(f"Загружаем документ: {filepath}")
        try:
            self.document = Document(filepath)
            self.logger.info(f"Документ загружен успешно: {filepath}")
        except Exception as e:
            self.logger.error(f"Ошибка при загрузке документа {filepath}: {e}")
            raise
        return self
    
    def process_document(self, input_filepath: str, prompt: str, llm_client=None) -> str:
        """
        Обрабатывает документ с помощью LLM: читает, отправляет в модель и сохраняет результат
        
        Args:
            input_filepath: Путь к исходному файлу (относительно docs/)
            prompt: Промт для LLM модели
            llm_client: Клиент LLM (OpenAIClient)
        
        Returns:
            Путь к обработанному файлу
        """
        # Формируем полный путь к исходному файлу
        if not input_filepath.startswith('docs/'):
            input_filepath = f"docs/{input_filepath}"
        
        if not os.path.exists(input_filepath):
            raise FileNotFoundError(f"Файл не найден: {input_filepath}")
        
        self.logger.info(f"Начинаем обработку документа: {input_filepath}")
        
        # Загружаем документ
        self.load_document(input_filepath)
        
        # Извлекаем текст из документа
        document_text = self._extract_text_from_document()
        self.logger.info(f"Извлечен текст из документа, длина: {len(document_text)} символов")
        
        # Отправляем в LLM
        if llm_client:
            self.logger.info("Отправляем текст в LLM")
            modified_text = self._process_with_llm(document_text, prompt, llm_client)
            if modified_text:
                self.logger.info("Применяем изменения от LLM к документу")
                self._apply_llm_changes(modified_text)
            else:
                self.logger.error("LLM не вернул результат")
                return None
        else:
            self.logger.error("LLM клиент не предоставлен")
            return None
        
        # Генерируем имя выходного файла: "new_" + старое имя
        base_name = os.path.basename(input_filepath)
        output_filename = f"new_{base_name}"
        
        # Убеждаемся, что папка new_docs существует
        os.makedirs("new_docs", exist_ok=True)
        output_path = f"new_docs/{output_filename}"
        
        # Сохраняем обработанный документ
        self.save_document(output_path)
        
        self.logger.info(f"Обработка документа завершена успешно: {output_path}")
        return output_path
    
    def _extract_text_from_document(self) -> str:
        """Извлекает весь текст из документа для отправки в LLM с уникальными маркерами"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        text_parts = []
        self.element_markers = []  # Сохраняем маркеры для последующего использования
        
        # 1. Извлекаем текст из параграфов основного содержимого
        for i, paragraph in enumerate(self.document.paragraphs):
            # Сохраняем даже пустые параграфы, так как они могут содержать форматирование
            marker = f"PARA_{i:04d}"
            self.element_markers.append(('paragraph', i, marker))
            text_parts.append(f"[{marker}] {paragraph.text.strip()}")
        
        # 2. Извлекаем текст из таблиц с улучшенной структурой
        table_idx = 0
        for table in self.document.tables:
            # Добавляем маркеры начала и конца таблицы для контекста
            table_start_marker = f"TABLE_START_{table_idx:02d}"
            self.element_markers.append(('table_start', table_idx, table_start_marker))
            text_parts.append(f"[{table_start_marker}]")
            
            # Обрабатываем каждую ячейку таблицы
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    marker = f"CELL_{table_idx:02d}_{row_idx:02d}_{cell_idx:02d}"
                    self.element_markers.append(('table_cell', table_idx, row_idx, cell_idx, marker))
                    
                    # Извлекаем текст из всех параграфов ячейки
                    cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    text_parts.append(f"[{marker}] {cell_text}")
            
            table_end_marker = f"TABLE_END_{table_idx:02d}"
            self.element_markers.append(('table_end', table_idx, table_end_marker))
            text_parts.append(f"[{table_end_marker}]")
            
            table_idx += 1
        
        # 3. Извлекаем текст из колонтитулов (первой секции документа)
        section = self.document.sections[0]
        
        # Обрабатываем верхние колонтитулы (headers)
        header_types = [
            (section.header, "HEADER"),
            (section.first_page_header, "HEADER_FIRST"),
            (section.even_page_header, "HEADER_EVEN")
        ]
        
        for header, header_type in header_types:
            if header is not None:
                for i, paragraph in enumerate(header.paragraphs):
                    marker = f"{header_type}_PARA_{i:04d}"
                    self.element_markers.append(('header', header_type, i, marker))
                    text_parts.append(f"[{marker}] {paragraph.text.strip()}")
                
                # Обрабатываем таблицы в колонтитулах
                for table_idx, table in enumerate(header.tables):
                    table_start_marker = f"{header_type}_TABLE_START_{table_idx:02d}"
                    self.element_markers.append(('header_table_start', header_type, table_idx, table_start_marker))
                    text_parts.append(f"[{table_start_marker}]")
                    
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            marker = f"{header_type}_CELL_{table_idx:02d}_{row_idx:02d}_{cell_idx:02d}"
                            self.element_markers.append(('header_table_cell', header_type, table_idx, row_idx, cell_idx, marker))
                            
                            cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                            text_parts.append(f"[{marker}] {cell_text}")
                    
                    table_end_marker = f"{header_type}_TABLE_END_{table_idx:02d}"
                    self.element_markers.append(('header_table_end', header_type, table_idx, table_end_marker))
                    text_parts.append(f"[{table_end_marker}]")
        
        # Обрабатываем нижние колонтитулы (footers)
        footer_types = [
            (section.footer, "FOOTER"),
            (section.first_page_footer, "FOOTER_FIRST"),
            (section.even_page_footer, "FOOTER_EVEN")
        ]
        
        for footer, footer_type in footer_types:
            if footer is not None:
                for i, paragraph in enumerate(footer.paragraphs):
                    marker = f"{footer_type}_PARA_{i:04d}"
                    self.element_markers.append(('footer', footer_type, i, marker))
                    text_parts.append(f"[{marker}] {paragraph.text.strip()}")
                
                # Обрабатываем таблицы в колонтитулах
                for table_idx, table in enumerate(footer.tables):
                    table_start_marker = f"{footer_type}_TABLE_START_{table_idx:02d}"
                    self.element_markers.append(('footer_table_start', footer_type, table_idx, table_start_marker))
                    text_parts.append(f"[{table_start_marker}]")
                    
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            marker = f"{footer_type}_CELL_{table_idx:02d}_{row_idx:02d}_{cell_idx:02d}"
                            self.element_markers.append(('footer_table_cell', footer_type, table_idx, row_idx, cell_idx, marker))
                            
                            cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                            text_parts.append(f"[{marker}] {cell_text}")
                    
                    table_end_marker = f"{footer_type}_TABLE_END_{table_idx:02d}"
                    self.element_markers.append(('footer_table_end', footer_type, table_idx, table_end_marker))
                    text_parts.append(f"[{table_end_marker}]")
        
        # Логируем информацию об извлеченных элементах
        para_count = len(self.document.paragraphs)
        table_count = len(self.document.tables)
        element_count = len(text_parts)
        
        self.logger.info(f"Извлечено элементов: {element_count} (параграфов: {para_count}, таблиц: {table_count})")
        
        return "\n".join(text_parts)
    
    
    def _process_with_llm(self, document_text: str, prompt: str, llm_client) -> str:
        """Отправляет текст в LLM и получает измененный результат"""
        try:
            self.logger.info("Начинаем обработку текста в LLM")
            
            # Подсчитываем количество строк в оригинальном тексте
            original_lines = document_text.split('\n')
            line_count = len(original_lines)
            
            # Формируем полный промт
            full_prompt = f"""
{prompt}

ВАЖНЕЙШИЕ ИНСТРУКЦИИ:
1. Ты получишь текст документа, где каждая строка начинается с уникального маркера в квадратных скобках, например [PARA_0005] или [CELL_02_01_03].
2. Ты ДОЛЖЕН сохранить ВСЕ эти маркеры в том же порядке и в том же месте в тексте. НЕ УДАЛЯЙ ИХ И НЕ ПЕРЕСТАВЛЯЙ.
3. Маркеры — это служебная информация для программы, которая позволяет ей сохранить исходное форматирование документа (шрифты, таблицы, отступы). Без них документ будет испорчен.
4. Внеси необходимые изменения в текст, который идет ПОСЛЕ маркера.
5. Если тебе нужно добавить новый абзац или элемент, НЕ ДЕЛАЙ ЭТОГО. Работай строго в рамках предоставленных маркеров.
6. Изменяй только текстовое содержимое. Не добавляй никаких комментариев, форматирования Markdown (вроде **жирный**) и не заключай ответ в ```.
7. Применяй указанные инструкции КО ВСЕМУ СТРОКАМ ТЕКСТА. НЕ ИГНОРИРУЙ НИКАКИЕ СТРОКИ.

ИСХОДНЫЙ ТЕКСТ С МАРКЕРАМИ:
{document_text}
"""
            
            # Подготавливаем сообщения для LLM
            messages = llm_client.prepare_messages(full_prompt)
            self.logger.info(f"Отправляем запрос в LLM, длина промта: {len(full_prompt)} символов")
            
            # Отправляем запрос
            result = llm_client.generate(messages)
            
            if result and result[0]:
                modified_text = result[0]
                self.logger.info(f"LLM вернул результат, длина: {len(modified_text)} символов")
                
                # Логируем количество строк для отладки
                lines = modified_text.split('\n')
                self.logger.info(f"LLM вернул {len(lines)} строк")
                
                return modified_text
            else:
                self.logger.error("LLM не вернул результат")
                return None
                
        except Exception as e:
            self.logger.error(f"Ошибка при работе с LLM: {e}")
            return None
    
    def _apply_llm_changes(self, modified_text: str):
        """Применяет изменения от LLM к документу с сохранением форматирования, используя маркеры."""
        if not self.document:
            raise ValueError("Документ не загружен!")

        self.logger.info(f"Применяем изменения от LLM, длина текста: {len(modified_text)} символов")

        # Создаем словарь: {маркер: новый_текст} для быстрого поиска
        changes_map = {}
        # Используем регулярное выражение для поиска маркеров и следующего за ними текста
        import re
        # pattern ищет [MARKER] и захватывает весь текст до следующего маркера или конца строки
        pattern = r'\[([A-Za-z_0-9]+)\]\s*(.*?)(?=\[[A-Za-z_0-9]+\]|$)'
        matches = re.findall(pattern, modified_text, re.DOTALL)

        for marker, new_content in matches:
            new_content = new_content.strip()
            if new_content:  # Игнорируем пустые совпадения
                changes_map[marker] = new_content
            else:
                # Если контент пустой, возможно, модель его удалила. Явно запишем пустую строку.
                changes_map[marker] = ""

        self.logger.info(f"Найдено и замаплено {len(changes_map)} маркеров в ответе LLM")

        # 1. Обновляем параграфы по маркерам
        paragraphs_updated = 0
        for elem_type, i, marker in [m for m in self.element_markers if m[0] == 'paragraph']:
            if marker in changes_map:
                new_text = changes_map[marker]
                self._update_paragraph_text(self.document.paragraphs[i], new_text)
                paragraphs_updated += 1

        # 2. Обновляем ячейки таблиц по маркерам
        tables_updated = 0
        for elem_type, table_idx, row_idx, cell_idx, marker in [m for m in self.element_markers if m[0] == 'table_cell']:
            if marker in changes_map:
                new_text = changes_map[marker]
                # Находим нужную ячейку и обновляем текст в ее первом параграфе (это упрощение, но работает для большинства случаев)
                target_cell = self.document.tables[table_idx].cell(row_idx, cell_idx)
                if target_cell.paragraphs:
                    self._update_paragraph_text(target_cell.paragraphs[0], new_text)
                tables_updated += 1

        self.logger.info(f"Изменения от LLM применены: параграфов: {paragraphs_updated}, ячеек таблиц: {tables_updated}")
    
    def apply_changes(self, changes: Dict[str, str]) -> 'WordProcessor':
        """Применяет изменения к документу с сохранением форматирования"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        self.logger.info("Применяем изменения")
        changes_made = 0
        
        # Изменяем параграфы
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                new_text = original_text
                
                # Применяем замены
                for old_text, new_text_replacement in changes.items():
                    if old_text in new_text:
                        new_text = new_text.replace(old_text, new_text_replacement)
                        self.logger.info(f"Заменяем: '{old_text}' -> '{new_text_replacement}'")
                        changes_made += 1
                
                # Обновляем текст с сохранением форматирования
                if new_text != original_text:
                    self._update_paragraph_text(paragraph, new_text)
        
        # Изменяем таблицы
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            original_text = paragraph.text
                            new_text = original_text
                            
                            # Применяем замены
                            for old_text, new_text_replacement in changes.items():
                                if old_text in new_text:
                                    new_text = new_text.replace(old_text, new_text_replacement)
                                    self.logger.info(f"В таблице заменяем: '{old_text}' -> '{new_text_replacement}'")
                                    changes_made += 1
                            
                            # Обновляем текст
                            if new_text != original_text:
                                self._update_paragraph_text(paragraph, new_text)
        
        self.logger.info(f"Изменения применены: {changes_made} замен")
        return self
    
    def _update_paragraph_text(self, paragraph, new_text: str):
        """Обновляет текст параграфа с сохранением форматирования"""
        try:
            if paragraph.runs:
                # Сохраняем форматирование из первого run
                first_run = paragraph.runs[0]
                original_font = first_run.font
                
                # Очищаем параграф
                paragraph.clear()
                
                # Добавляем новый текст с сохранением форматирования
                new_run = paragraph.add_run(new_text)
                
                # Копируем форматирование шрифта
                if original_font.name:
                    new_run.font.name = original_font.name
                if original_font.size:
                    new_run.font.size = original_font.size
                if original_font.bold is not None:
                    new_run.font.bold = original_font.bold
                if original_font.italic is not None:
                    new_run.font.italic = original_font.italic
                if original_font.underline is not None:
                    new_run.font.underline = original_font.underline
                if original_font.color.rgb:
                    new_run.font.color.rgb = original_font.color.rgb
            else:
                # Если нет runs, просто заменяем текст
                paragraph.text = new_text
        except Exception as e:
            self.logger.error(f"Ошибка при обновлении параграфа: {e}")
            # Fallback - просто заменяем текст
            paragraph.text = new_text
    
    def save_document(self, output_path: str) -> 'WordProcessor':
        """Сохраняет документ"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        self.logger.info(f"Сохраняем документ: {output_path}")
        try:
            self.document.save(output_path)
            self.logger.info(f"Документ сохранен успешно: {output_path}")
        except Exception as e:
            self.logger.error(f"Ошибка при сохранении документа {output_path}: {e}")
            raise
        return self
    
    def show_document_info(self):
        """Показывает информацию о документе"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        self.logger.info("Информация о документе:")
        self.logger.info("-" * 40)
        
        # Подсчитываем элементы
        paragraphs = [p for p in self.document.paragraphs if p.text.strip()]
        tables = self.document.tables
        
        self.logger.info(f"Параграфов: {len(paragraphs)}")
        self.logger.info(f"Таблиц: {len(tables)}")
        
        # Показываем форматирование
        self.logger.info("Форматирование:")
        for i, paragraph in enumerate(paragraphs[:3]):  # Показываем первые 3
            self.logger.info(f"Параграф {i+1}: {paragraph.text[:50]}...")
            if paragraph.runs:
                for j, run in enumerate(paragraph.runs[:2]):  # Показываем первые 2 runs
                    if run.text.strip():
                        self.logger.info(f"  Run {j+1}: жирный={run.bold}, курсив={run.italic}")