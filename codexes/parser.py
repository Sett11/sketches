#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Парсер DOCX файлов
"""

import json
import os
import re
from typing import Dict, List, Any, Optional
from docx import Document
from mylogger import Logger


class Parser:
    """Парсер DOCX файлов"""
    
    def __init__(self):
        # Инициализация логгера
        self.logger = Logger('Parser', 'logs/parser.log')
        self.logger.info("Инициализация DOCX парсера")

    def read_docx_file(self, file_path: str) -> Document:
        """Чтение DOCX файла"""
        self.logger.debug(f"Чтение DOCX файла: {file_path}")
        try:
            doc = Document(file_path)
            self.logger.info(f"DOCX файл успешно прочитан: {file_path}")
            return doc
        except Exception as e:
            self.logger.error(f"Ошибка чтения DOCX файла {file_path}: {str(e)}")
            raise

    def clean_title(self, title: str) -> str:
        """Очистка названия от лишних символов"""
        # Убираем точки в начале и конце
        title = title.strip(' .')
        # Убираем лишние пробелы
        title = re.sub(r'\s+', ' ', title)
        return title

    def extract_structure(self, doc: Document) -> Dict[str, Any]:
        """Извлечение структуры документа из DOCX"""
        self.logger.debug("Начало извлечения структуры из DOCX (иерархический режим)")
        
        structure = {
            'document': '',
            'parts': []
        }
        
        # Новый основной алгоритм: иерархический split
        full_text = self.get_full_text(doc)
        self.logger.debug(f"Длина полного текста: {len(full_text)}")

        parts = self.split_level(full_text, level='part')
        self.logger.info(f"Найдено частей: {len(parts)}")
        if not parts:
            parts = [{
                'number': '1',
                'title': 'Основная часть',
                'body': full_text
            }]

        for p_idx, part in enumerate(parts, 1):
            part_obj = {
                'number': part['number'],
                'title': self.clean_title(part['title']),
                'sections': []
            }
            
            sections = self.split_level(part['body'], level='section')
            if not sections:
                sections = [{
                    'number': '1',
                    'title': 'Общий раздел',
                    'body': part['body']
                }]

            for s_idx, section in enumerate(sections, 1):
                section_obj = {
                    'number': section['number'],
                    'title': self.clean_title(section['title']),
                        'chapters': []
                    }

                chapters = self.split_level(section['body'], level='chapter')
                if not chapters:
                    chapters = [{
                        'number': '1',
                        'title': 'Общая глава',
                        'body': section['body']
                    }]

                for c_idx, chapter in enumerate(chapters, 1):
                    chapter_number = self.normalize_chapter_article_number(chapter['number'])
                    chapter_obj = {
                        'number': chapter_number,
                        'title': self.clean_title(chapter['title']),
                        'articles': []
                    }

                    articles = self.split_level(chapter['body'], level='article')

                    for a_idx, article in enumerate(articles, 1):
                        article_number = self.normalize_chapter_article_number(article['number'])
                        article_obj = {
                            'number': article_number,
                            'title': self.clean_title(article['title']),
                            'notes': [],
                            'intro': '',
                        'paragraphs': []
                    }
                
                        intro_lines: List[str] = []
                        started_numbered = False
                        for line in article['body'].split('\n'):
                            text_line = line.strip()
                            if not text_line:
                                continue
                            # собираем примечания в отдельный массив
                            if re.match(r'^\s*Примечан[ие][\w\W]*', text_line, flags=re.IGNORECASE):
                                article_obj['notes'].append(text_line)
                                continue
                            if not started_numbered and not self.is_paragraph(text_line):
                                intro_lines.append(text_line)
                                continue
                            if self.is_paragraph(text_line):
                                started_numbered = True
                                paragraph_data = self.extract_paragraph_structure(text_line)
                                if paragraph_data:
                                    article_obj['paragraphs'].append(paragraph_data)
                        if intro_lines:
                            article_obj['intro'] = "\n".join(intro_lines)

                        chapter_obj['articles'].append(article_obj)

                    section_obj['chapters'].append(chapter_obj)

                part_obj['sections'].append(section_obj)

            structure['parts'].append(part_obj)

        self.logger.info("Иерархическое извлечение структуры завершено")
        return structure

    def is_paragraph(self, text: str) -> bool:
        """Проверка, является ли текст пунктом или подпунктом"""
        patterns = [
            r'^(\d+)\.\s',              # 1. ...
            r'^(?:\d+\.){2,}\s',        # 1.2.3. ...
            r'^(\d+)\.(\d+)\.\s',     # 1.1. ...
            r'^(\d+)\)\s',              # 1) ...
            r'^(\d+(?:\.\d+)+)\)\s',  # 1.1) ...
            r'^([A-Za-zА-Яа-яЁё])\)\s',  # a) б) ...
            r'^[\-—•]\s',                # -, —, • ...
        ]
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False

    def extract_paragraph_structure(self, text: str) -> Dict[str, Any]:
        """Извлечение структуры пункта (включая подпункты)"""
        m = re.match(r'^(?P<num>(?:\d+\.){1,}\d*|\d+(?:\.\d+)+|\d+\)|[A-Za-zА-Яа-яЁё]\)|[\-—•])\s+(?P<text>.+)', text)
        if m:
            return {
                'number': m.group('num'),
                'text': m.group('text').strip(),
                'subparagraphs': []
            }
        return None

    def get_full_text(self, doc: Document) -> str:
        """Собирает текст документа, объединяя непустые параграфы построчно."""
        lines: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or '').strip()
            if t:
                lines.append(t)
        text = "\n".join(lines)
        text = re.sub(r"\r\n?", "\n", text)
        # Убираем лишние пробелы в начале/конце строк
        text = "\n".join(s.strip() for s in text.split("\n"))
        return text

    def split_level(self, text: str, level: str) -> List[Dict[str, Any]]:
        """Разделяет текст по заголовкам уровня (part/section/chapter/article) и возвращает список блоков."""
        patterns = {
            # Разрешаем ведущие пробелы и точку после номера; расширяем главы римскими цифрами
            'part':    r'(?m)^\s*(?P<kw>ЧАСТЬ)\s+(?P<num>[IVXLC]+|[А-ЯЁ]+|[0-9]+)\.?\s*(?P<title>.*)$',
            'section': r'(?m)^\s*(?P<kw>Раздел)\s+(?P<num>[IVXLC]+|[А-ЯЁ]+|[0-9]+)\.?\s*(?P<title>.*)$',
            'chapter': r'(?m)^\s*(?P<kw>Глава)\s+(?P<num>[IVXLC]+|[0-9]+)\.?\s*(?P<title>.*)$',
            # Для статей допускаем перенос между словом и номером
            'article': r'(?m)^\s*(?P<kw>Статья)\s*(?:\n\s*)?(?P<num>[0-9]+)\.?\s*(?P<title>.*)$',
        }
        if level not in patterns:
            return []
        regex = re.compile(patterns[level], re.IGNORECASE)
        matches = list(regex.finditer(text))
        if not matches:
            return []

        items: List[Dict[str, Any]] = []
        for idx, m in enumerate(matches):
            end_heading = m.end()
            next_start = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
            body = text[end_heading:next_start]
            title_raw = (m.group('title') or '').strip()
            # если "заголовок" фактически начинается с маркера пункта — считаем его пустым
            if title_raw and self.is_paragraph(title_raw):
                title_raw = ''
            # Если после ключевого слова пусто/только пунктуация — взять следующую строку как заголовок
            if not title_raw or re.fullmatch(r'[.·•—\-]*', title_raw):
                title_raw = self.peek_next_nonempty_line(text, end_heading, avoid_paragraphs=True, avoid_headings=True)
            item = {
                'number': (m.group('num') or '').strip(),
                'title': self.clean_title(title_raw or ''),
                'body': body
            }
            if self.is_amendment_or_note(item['title']):
                self.logger.debug(f"Отфильтрован заголовок {level}: {item['title']}")
                continue
            items.append(item)
        return items

    def peek_next_nonempty_line(self, text: str, pos: int, avoid_paragraphs: bool = False, avoid_headings: bool = False) -> str:
        tail = text[pos:]
        for line in tail.split('\n'):
            s = line.strip()
            if s:
                if avoid_paragraphs and self.is_paragraph(s):
                    continue
                if avoid_headings and self.is_heading_line(s):
                    continue
                return s
        return ''

    def is_heading_line(self, text: str) -> bool:
        return bool(re.match(r'^(ЧАСТЬ|Раздел|Глава|Статья)\b', text, re.IGNORECASE))

    def normalize_chapter_article_number(self, value: Any) -> int:
        try:
            if isinstance(value, int):
                return value
            s = str(value).strip()
            return int(s) if s.isdigit() else 0
        except Exception:
            return 0

    def is_amendment_or_note(self, text: str) -> bool:
        if not text:
            return False
        patterns = [
            r'в ред\.',
            r'федерального закона',
            r'постановлени[ея]',
            r'конституционного суда',
            r'примечани[ея]'
        ]
        return any(re.search(p, text, re.IGNORECASE) for p in patterns)

    def extract_number_after_keyword(self, text: str, keyword: str) -> str:
        """Извлечение номера после ключевого слова"""
        pattern = rf'{keyword}\s+([IVX]+|[А-ЯЁ]+|[0-9]+)'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)
        return ""

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Парсинг одного DOCX файла"""
        self.logger.info(f"Начало обработки файла: {file_path}")
        try:
            # Читаем файл
            doc = self.read_docx_file(file_path)
            self.logger.debug(f"DOCX файл прочитан, количество параграфов: {len(doc.paragraphs)}")
            
            # Извлекаем структуру
            structure = self.extract_structure(doc)
            
            # Устанавливаем название документа из имени файла
            file_name = os.path.basename(file_path)
            structure['document'] = file_name.replace('.docx', '')
            
            self.logger.info(f"Файл {file_name} успешно обработан")
            return structure
            
        except Exception as e:
            self.logger.error(f"Ошибка при обработке файла {file_path}: {str(e)}")
            return {
                'document': os.path.basename(file_path).replace('.docx', ''),
                'error': f'Ошибка обработки файла: {str(e)}',
                'parts': []
            }

    def process_all_files(self, input_dir: str, output_dir: str):
        """Обработка всех DOCX файлов в директории"""
        self.logger.info(f"Начало обработки всех файлов из {input_dir} в {output_dir}")
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            self.logger.info(f"Создана директория {output_dir}")
        
        docx_files = [f for f in os.listdir(input_dir) if f.endswith('.docx')]
        self.logger.info(f"Найдено {len(docx_files)} DOCX файлов для обработки")
        
        for i, docx_file in enumerate(docx_files, 1):
            self.logger.info(f"Обрабатываю файл {i}/{len(docx_files)}: {docx_file}")
            print(f"Обрабатываю файл {i}/{len(docx_files)}: {docx_file}")
            
            input_path = os.path.join(input_dir, docx_file)
            output_path = os.path.join(output_dir, docx_file.replace('.docx', '.json'))
            
            # Парсим файл
            structure = self.parse_file(input_path)
            
            # Логируем результат парсинга
            parts_count = len(structure.get('parts', []))
            if 'error' in structure:
                self.logger.error(f"Ошибка при обработке {docx_file}: {structure['error']}")
            else:
                self.logger.info(f"Файл {docx_file} обработан успешно. Найдено частей: {parts_count}")
            
            # Сохраняем JSON
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(structure, f, ensure_ascii=False, indent=2)
            
            self.logger.info(f"Сохранен: {output_path}")
            print(f"Сохранен: {output_path}")
        
        self.logger.info("Обработка всех файлов завершена")


def main():
    """Основная функция"""
    parser = Parser()
    parser.process_all_files('docs', 'jsons')


if __name__ == "__main__":
    main()
