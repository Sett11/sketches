#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Отладочный скрипт для проверки содержимого .docx файлов
"""

from docx import Document
import re

def debug_docx_file(file_path):
    """Отлаживает содержимое .docx файла"""
    print(f"Анализируем файл: {file_path}")
    
    try:
        doc = Document(file_path)
        print(f"Всего абзацев: {len(doc.paragraphs)}")
        
        # Показываем первые 10 абзацев
        print("\nПервые 10 абзацев:")
        for i, paragraph in enumerate(doc.paragraphs[:10]):
            if paragraph.text.strip():
                print(f"{i}: {paragraph.text[:150]}...")
        
        # Ищем статьи с разными паттернами
        print("\nПоиск статей:")
        article_count = 0
        patterns = [
            r'\d+\s*[Сс]татья',
            r'[Сс]татья\s*\d+',
            r'Статья\s*\d+',
            r'статья\s*\d+'
        ]
        
        for pattern in patterns:
            print(f"\nПаттерн: {pattern}")
            count = 0
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if re.search(pattern, text):
                    print(f"Найдена статья в абзаце {i}: {text[:100]}...")
                    count += 1
                    if count >= 3:  # Показываем только первые 3 для каждого паттерна
                        break
            print(f"Найдено с паттерном '{pattern}': {count}")
            article_count += count
        
        print(f"\nВсего найдено статей: {article_count}")
        
        # Показываем больше абзацев для анализа
        print("\nАбзацы 10-20:")
        for i, paragraph in enumerate(doc.paragraphs[10:20]):
            if paragraph.text.strip():
                print(f"{i+10}: {paragraph.text[:150]}...")
        
    except Exception as e:
        print(f"Ошибка: {e}")

if __name__ == "__main__":
    # Тестируем первый файл
    debug_docx_file("docs/1. Гражданский-кодекс-Российской-Федерации-_часть-первая_-от-30.11.94-N-51-ФЗ.docx")
