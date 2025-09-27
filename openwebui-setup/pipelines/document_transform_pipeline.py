#!/usr/bin/env python3
"""
Единый Pipeline для трансформации Word документов с помощью LLM
Содержит всю логику: OpenRouterClient, WordProcessor, DocumentTransformPipeline
"""

import os
import re
from typing import Dict, List, Any, Optional
from datetime import datetime
import requests
from docx import Document

# Импортируем наш логгер
import sys
# Добавляем родительскую директорию в путь для импорта
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)
from mylogger import Logger

# Настройка логирования
logger = Logger('DOCUMENT_TRANSFORM', 'logs/document_transform.log')

class OpenRouterClient:
    """Клиент для работы с OpenRouter API"""
    
    def __init__(self, model_name: str = None, api_key: str = None, site_url: str = None, site_name: str = None):
        # Получаем значения из переменных окружения
        self.api_key = api_key or os.getenv('OPENROUTER_API_KEY')
        self.model_name = model_name or os.getenv('MODEL_NAME', 'google/gemma-3-27b-it:free')
        self.site_url = site_url or os.getenv('SITE_URL', '')
        self.site_name = site_name or os.getenv('SITE_NAME', '')
        
        if not self.api_key:
            raise ValueError("OPENROUTER_API_KEY не найден в переменных окружения")
        
        self.base_url = "https://openrouter.ai/api/v1/chat/completions"
        logger.info(f"OpenRouterClient инициализирован: model={self.model_name}")

    def generate(self, messages: list, max_retries: int = 3, delay: int = 600, 
                 temperature: float = 0.3, max_tokens: int = 2048, idop: int = 0):
        """Вызывает OpenRouter API с обработкой исключений"""
        retries = 0
        
        while retries < max_retries:
            try:
                headers = {
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json"
                }
                
                if self.site_url:
                    headers["Referer"] = self.site_url
                if self.site_name:
                    headers["X-Title"] = self.site_name
                
                data = {
                    "model": self.model_name,
                    "messages": messages,
                    "temperature": temperature,
                    "max_tokens": max_tokens
                }
                
                response = requests.post(self.base_url, headers=headers, json=data)
                response.raise_for_status()
                
                result = response.json()
                
                if 'choices' in result and len(result['choices']) > 0:
                    content = result['choices'][0]['message']['content']
                    usage = result.get('usage', {})
                    prompt_tokens = usage.get('prompt_tokens', 0)
                    completion_tokens = usage.get('completion_tokens', 0)
                    
                    if idop != 0:
                        logger.info(f"LLM call - idop: {idop}, model: {self.model_name}, prompt_tokens: {int(prompt_tokens)}, completion_tokens: {int(completion_tokens)}")
                    else:
                        logger.info("Call with NULL idop")
                    
                    return (content, prompt_tokens, completion_tokens)
                else:
                    logger.error("No choices in OpenRouter response")
                    return None
            
            except requests.exceptions.RequestException as e:
                logger.error(f"Request Error: {e}")
                if retries < max_retries - 1:
                    logger.info(f"Retrying in {delay} seconds...")
                    retries += 1
                    time.sleep(delay)
                else:
                    return None
            
            except Exception as e:
                logger.error(f"Unexpected Error: {e}")
                return None
        
        logger.error("Max retries reached. Failed to get response.")
        return None

    @staticmethod
    def prepare_messages(prompt: str, system_message: str = "") -> list:
        """Формирует список сообщений для OpenRouter API"""
        messages = []
        if system_message:
            messages.append({"role": "system", "content": system_message})
        messages.append({"role": "user", "content": prompt})
        return messages


class WordProcessor:
    """Процессор Word документов с сохранением форматирования"""
    
    def __init__(self):
        self.document = None
        self.logger = Logger('WordProcessor', 'logs/word_processor.log')
        self.logger.info("WordProcessor инициализирован")
    
    def load_document(self, filepath: str) -> 'WordProcessor':
        """Загружает документ из файла"""
        self.logger.info(f"Загружаем документ: {filepath}")
        try:
            self.document = Document(filepath)
            self.logger.info(f"Документ загружен успешно: {filepath}")
        except Exception as e:
            self.logger.error(f"Ошибка при загрузке документа {filepath}: {e}")
            raise
        return self
    
    def process_document(self, input_filepath: str, prompt: str, llm_client=None) -> str:
        """Обрабатывает документ с помощью LLM"""
        if not os.path.exists(input_filepath):
            raise FileNotFoundError(f"Файл не найден: {input_filepath}")
        
        self.logger.info(f"Начинаем обработку документа: {input_filepath}")
        
        # Загружаем документ
        self.load_document(input_filepath)
        
        # Извлекаем текст из документа
        document_text = self._extract_text_from_document()
        self.logger.info(f"Извлечен текст из документа, длина: {len(document_text)} символов")
        
        # Отправляем в LLM
        if llm_client:
            self.logger.info("Отправляем текст в LLM")
            modified_text = self._process_with_llm(document_text, prompt, llm_client)
            if modified_text:
                self.logger.info("Применяем изменения от LLM к документу")
                self._apply_llm_changes(modified_text)
            else:
                self.logger.error("LLM не вернул результат")
                return None
        else:
            self.logger.error("LLM клиент не предоставлен")
            return None
        
        # Генерируем имя выходного файла
        base_name = os.path.basename(input_filepath)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"transformed_{timestamp}_{base_name}"
        
        # Создаем папку new_docs если не существует
        os.makedirs("new_docs", exist_ok=True)
        output_path = f"new_docs/{output_filename}"
        
        # Сохраняем обработанный документ
        self.save_document(output_path)
        
        self.logger.info(f"Обработка документа завершена успешно: {output_path}")
        return output_path
    
    def _extract_text_from_document(self) -> str:
        """Извлекает весь текст из документа для отправки в LLM с уникальными маркерами"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        text_parts = []
        self.element_markers = []  # Сохраняем маркеры для последующего использования
        
        # 1. Извлекаем текст из параграфов основного содержимого
        for i, paragraph in enumerate(self.document.paragraphs):
            marker = f"PARA_{i:04d}"
            self.element_markers.append(('paragraph', i, marker))
            text_parts.append(f"[{marker}] {paragraph.text.strip()}")
        
        # 2. Извлекаем текст из таблиц
        table_idx = 0
        for table in self.document.tables:
            table_start_marker = f"TABLE_START_{table_idx:02d}"
            self.element_markers.append(('table_start', table_idx, table_start_marker))
            text_parts.append(f"[{table_start_marker}]")
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    marker = f"CELL_{table_idx:02d}_{row_idx:02d}_{cell_idx:02d}"
                    self.element_markers.append(('table_cell', table_idx, row_idx, cell_idx, marker))
                    
                    cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    text_parts.append(f"[{marker}] {cell_text}")
            
            table_end_marker = f"TABLE_END_{table_idx:02d}"
            self.element_markers.append(('table_end', table_idx, table_end_marker))
            text_parts.append(f"[{table_end_marker}]")
            
            table_idx += 1
        
        # Логируем информацию об извлеченных элементах
        para_count = len(self.document.paragraphs)
        table_count = len(self.document.tables)
        element_count = len(text_parts)
        
        self.logger.info(f"Извлечено элементов: {element_count} (параграфов: {para_count}, таблиц: {table_count})")
        
        return "\n".join(text_parts)
    
    def _process_with_llm(self, document_text: str, prompt: str, llm_client) -> str:
        """Отправляет текст в LLM и получает измененный результат"""
        try:
            self.logger.info("Начинаем обработку текста в LLM")
            
            # Формируем полный промт
            full_prompt = f"""
{prompt}

ВАЖНЕЙШИЕ ИНСТРУКЦИИ:
1. Ты получишь текст документа, где каждая строка начинается с уникального маркера в квадратных скобках, например [PARA_0005] или [CELL_02_01_03].
2. Ты ДОЛЖЕН сохранить ВСЕ эти маркеры в том же порядке и в том же месте в тексте. НЕ УДАЛЯЙ ИХ И НЕ ПЕРЕСТАВЛЯЙ.
3. Маркеры — это служебная информация для программы, которая позволяет ей сохранить исходное форматирование документа (шрифты, таблицы, отступы). Без них документ будет испорчен.
4. Внеси необходимые изменения в текст, который идет ПОСЛЕ маркера.
5. Если тебе нужно добавить новый абзац или элемент, НЕ ДЕЛАЙ ЭТОГО. Работай строго в рамках предоставленных маркеров.
6. Изменяй только текстовое содержимое. Не добавляй никаких комментариев, форматирования Markdown (вроде **жирный**) и не заключай ответ в ```.
7. Применяй указанные инструкции КО ВСЕМУ СТРОКАМ ТЕКСТА. НЕ ИГНОРИРУЙ НИКАКИЕ СТРОКИ.
8. Если в тексте есть таблицы, то не смешивай текст, который вокруг таблицы, с текстом, который внутри таблицы.

ИСХОДНЫЙ ТЕКСТ С МАРКЕРАМИ:
{document_text}
"""
            
            # Подготавливаем сообщения для LLM
            messages = llm_client.prepare_messages(full_prompt)
            self.logger.info(f"Отправляем запрос в LLM, длина промта: {len(full_prompt)} символов")
            
            # Отправляем запрос
            result = llm_client.generate(messages)
            
            if result and result[0]:
                modified_text = result[0]
                self.logger.info(f"LLM вернул результат, длина: {len(modified_text)} символов")
                
                lines = modified_text.split('\n')
                self.logger.info(f"LLM вернул {len(lines)} строк")
                
                return modified_text
            else:
                self.logger.error("LLM не вернул результат")
                return None
                
        except Exception as e:
            self.logger.error(f"Ошибка при работе с LLM: {e}")
            return None
    
    def _apply_llm_changes(self, modified_text: str):
        """Применяет изменения от LLM к документу с сохранением форматирования"""
        if not self.document:
            raise ValueError("Документ не загружен!")

        self.logger.info(f"Применяем изменения от LLM, длина текста: {len(modified_text)} символов")

        # Создаем словарь: {маркер: новый_текст} для быстрого поиска
        changes_map = {}
        pattern = r'\[([A-Za-z_0-9]+)\]\s*(.*?)(?=\[[A-Za-z_0-9]+\]|$)'
        matches = re.findall(pattern, modified_text, re.DOTALL)

        for marker, new_content in matches:
            new_content = new_content.strip()
            if new_content:
                changes_map[marker] = new_content
            else:
                changes_map[marker] = ""

        self.logger.info(f"Найдено и замаплено {len(changes_map)} маркеров в ответе LLM")

        # Обновляем параграфы по маркерам
        paragraphs_updated = 0
        for elem_type, i, marker in [m for m in self.element_markers if m[0] == 'paragraph']:
            if marker in changes_map:
                new_text = changes_map[marker]
                self._update_paragraph_text(self.document.paragraphs[i], new_text)
                paragraphs_updated += 1

        # Обновляем ячейки таблиц по маркерам
        tables_updated = 0
        for elem_type, table_idx, row_idx, cell_idx, marker in [m for m in self.element_markers if m[0] == 'table_cell']:
            if marker in changes_map:
                new_text = changes_map[marker]
                target_cell = self.document.tables[table_idx].cell(row_idx, cell_idx)
                if target_cell.paragraphs:
                    self._update_paragraph_text(target_cell.paragraphs[0], new_text)
                tables_updated += 1

        self.logger.info(f"Изменения от LLM применены: параграфов: {paragraphs_updated}, ячеек таблиц: {tables_updated}")
    
    def _update_paragraph_text(self, paragraph, new_text: str):
        """Обновляет текст параграфа с сохранением форматирования"""
        try:
            if paragraph.runs:
                # Сохраняем форматирование из первого run
                first_run = paragraph.runs[0]
                original_font = first_run.font
                
                # Очищаем параграф
                paragraph.clear()
                
                # Добавляем новый текст с сохранением форматирования
                new_run = paragraph.add_run(new_text)
                
                # Копируем форматирование шрифта
                if original_font.name:
                    new_run.font.name = original_font.name
                if original_font.size:
                    new_run.font.size = original_font.size
                if original_font.bold is not None:
                    new_run.font.bold = original_font.bold
                if original_font.italic is not None:
                    new_run.font.italic = original_font.italic
                if original_font.underline is not None:
                    new_run.font.underline = original_font.underline
                if original_font.color.rgb:
                    new_run.font.color.rgb = original_font.color.rgb
            else:
                # Если нет runs, просто заменяем текст
                paragraph.text = new_text
        except Exception as e:
            self.logger.error(f"Ошибка при обновлении параграфа: {e}")
            # Fallback - просто заменяем текст
            paragraph.text = new_text
    
    def save_document(self, output_path: str) -> 'WordProcessor':
        """Сохраняет документ"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        self.logger.info(f"Сохраняем документ: {output_path}")
        try:
            self.document.save(output_path)
            self.logger.info(f"Документ сохранен успешно: {output_path}")
        except Exception as e:
            self.logger.error(f"Ошибка при сохранении документа {output_path}: {e}")
            raise
        return self


class DocumentTransformPipeline:
    """Основной pipeline для трансформации документов"""
    
    def __init__(self):
        self.logger = Logger('DOCUMENT_TRANSFORM_PIPELINE', 'logs/document_transform_pipeline.log')
        self.llm_client = None
        self.word_processor = None
        self.setup_llm_client()
        self.setup_word_processor()
        self.logger.info("DocumentTransformPipeline инициализирован")
    
    def setup_llm_client(self):
        """Инициализация LLM клиента"""
        try:
            self.llm_client = OpenRouterClient()
            self.logger.info("✅ LLM клиент инициализирован")
        except Exception as e:
            self.logger.error(f"❌ Ошибка инициализации LLM клиента: {e}")
            self.llm_client = None
    
    def setup_word_processor(self):
        """Инициализация процессора Word документов"""
        try:
            self.word_processor = WordProcessor()
            self.logger.info("✅ Word процессор инициализирован")
        except Exception as e:
            self.logger.error(f"❌ Ошибка инициализации Word процессора: {e}")
            self.word_processor = None
    
    def transform_document(self, file_path: str, prompt: str) -> Dict[str, Any]:
        """Основная функция трансформации документа"""
        try:
            self.logger.info(f"Начинаем трансформацию документа: {file_path}")
            
            if not self.llm_client or not self.word_processor:
                return {
                    "success": False,
                    "error": "LLM клиент или Word процессор не инициализированы"
                }
            
            # Проверяем существование файла
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Файл не найден: {file_path}"
                }
            
            # Обрабатываем документ
            result_path = self.word_processor.process_document(file_path, prompt, self.llm_client)
            
            if result_path and os.path.exists(result_path):
                return {
                    "success": True,
                    "original_file": file_path,
                    "transformed_file": result_path,
                    "message": "Документ успешно трансформирован"
                }
            else:
                return {
                    "success": False,
                    "error": "Ошибка при обработке документа"
                }
                
        except Exception as e:
            self.logger.error(f"Ошибка трансформации документа: {e}")
            return {
                "success": False,
                "error": str(e)
            }


# Создание экземпляра Pipeline (ленивая инициализация)
pipeline = None

def get_pipeline():
    """Получение экземпляра pipeline с ленивой инициализацией"""
    global pipeline
    if pipeline is None:
        try:
            pipeline = DocumentTransformPipeline()
        except Exception as e:
            logger.error(f"❌ Ошибка инициализации pipeline: {e}")
            return None
    return pipeline

# Старые функции pipe(), on_startup(), on_shutdown() удалены
# Теперь используется только get_pipeline() для ленивой инициализации